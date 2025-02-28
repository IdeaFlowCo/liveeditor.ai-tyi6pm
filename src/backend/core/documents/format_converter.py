"""
Provides document format conversion capabilities for the AI writing enhancement platform,
enabling import and export between different document formats like Plain Text, Rich Text,
HTML, Markdown, Microsoft Word, and PDF. It handles automatic format detection, conversion
between formats, and maintains document structure integrity during conversion.
"""

import io
import os
import typing
import re
from typing import Union, Dict, List, Optional, Callable, Tuple, Any

# Third-party libraries for format conversion
import html2text  # html2text 2020.1.16
import markdown  # markdown 3.4.1
import docx  # python-docx 0.8.11
import weasyprint  # weasyprint 57.1
import striprtf  # striprtf 0.0.22
import mammoth  # mammoth 1.5.1

# Internal imports
from ..utils.logger import get_logger
from ..utils.validators import is_valid_document_format, ALLOWED_DOCUMENT_FORMATS

# Initialize logger
logger = get_logger(__name__)

# MIME type constants
HTML_MIME_TYPE = "text/html"
PLAIN_TEXT_MIME_TYPE = "text/plain"
MARKDOWN_MIME_TYPE = "text/markdown"
DOCX_MIME_TYPE = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
RTF_MIME_TYPE = "application/rtf"
PDF_MIME_TYPE = "application/pdf"

# Format to MIME type mapping
FORMAT_MIME_TYPES = {
    "html": HTML_MIME_TYPE,
    "txt": PLAIN_TEXT_MIME_TYPE,
    "md": MARKDOWN_MIME_TYPE,
    "docx": DOCX_MIME_TYPE,
    "rtf": RTF_MIME_TYPE,
    "pdf": PDF_MIME_TYPE
}

# Default format if detection fails
DEFAULT_FORMAT = "html"


class FormatConversionError(Exception):
    """Exception raised when document format conversion fails."""
    
    def __init__(self, message: str, source_format: str, target_format: str, original_error: Exception = None):
        """
        Initialize FormatConversionError with detailed information.
        
        Args:
            message: Error message
            source_format: Source document format
            target_format: Target document format
            original_error: Original exception that caused the error
        """
        super().__init__(message)
        self.source_format = source_format
        self.target_format = target_format
        self.original_error = original_error


class UnsupportedFormatError(Exception):
    """Exception raised when an unsupported document format is requested."""
    
    def __init__(self, format_name: str, operation: str, message: str = None):
        """
        Initialize UnsupportedFormatError.
        
        Args:
            format_name: Name of the unsupported format
            operation: Operation being performed (import/export)
            message: Error message
        """
        if not message:
            message = f"Unsupported format '{format_name}' for {operation} operation"
        super().__init__(message)
        self.format_name = format_name
        self.operation = operation


class FormatConverter:
    """Main class that handles document format conversion between supported formats."""
    
    def __init__(self):
        """
        Initializes the format converter with supported formats and conversion mappings.
        """
        # Define supported import and export formats
        self._import_formats = ["html", "txt", "md", "docx", "rtf"]
        self._export_formats = ["html", "txt", "md", "docx", "pdf"]
        
        # Initialize converters dictionary
        # Key format: "source_format:target_format"
        self._converters = {}
        
        # Initialize format detectors
        self._detectors = {
            "html": self._detect_html,
            "md": self._detect_markdown,
            "docx": self._detect_docx,
            "rtf": self._detect_rtf
        }
        
        # Register built-in converters
        self.register_converter("html", "md", self._convert_html_to_markdown)
        self.register_converter("html", "txt", self._convert_html_to_text)
        self.register_converter("md", "html", self._convert_markdown_to_html)
        self.register_converter("md", "txt", self._convert_markdown_to_text)
        self.register_converter("txt", "html", self._convert_text_to_html)
        self.register_converter("docx", "html", self._convert_docx_to_html)
        self.register_converter("rtf", "txt", self._convert_rtf_to_text)
        self.register_converter("html", "pdf", self._convert_html_to_pdf)
        self.register_converter("html", "docx", self._convert_html_to_docx)
    
    def detect_format(self, content: Union[str, bytes], filename: str = None, mime_type: str = None) -> str:
        """
        Detects the format of document content based on content analysis.
        
        Args:
            content: Document content to analyze
            filename: Original filename, if available
            mime_type: MIME type, if available
            
        Returns:
            Detected format as a string (extension without dot)
        """
        # Check if filename has a recognized extension
        if filename:
            try:
                ext = os.path.splitext(filename)[1].lower()
                if ext.startswith('.'):
                    ext = ext[1:]  # Remove leading dot
                if is_valid_document_format(ext):
                    logger.debug(f"Format detected from filename: {ext}")
                    return ext
            except Exception as e:
                logger.warning(f"Error extracting extension from filename: {str(e)}")
        
        # Check MIME type if provided
        if mime_type:
            for fmt, mime in FORMAT_MIME_TYPES.items():
                if mime.lower() == mime_type.lower():
                    logger.debug(f"Format detected from MIME type: {fmt}")
                    return fmt
        
        # Analyze content using format detectors
        for fmt, detector in self._detectors.items():
            try:
                if detector(content):
                    logger.debug(f"Format detected by content analysis: {fmt}")
                    return fmt
            except Exception as e:
                logger.warning(f"Error in format detector for {fmt}: {str(e)}")
        
        # If no format detected, return default
        logger.info(f"Could not detect format, using default: {DEFAULT_FORMAT}")
        return DEFAULT_FORMAT
    
    def convert(self, content: Union[str, bytes], source_format: str, target_format: str, 
                options: Dict[str, Any] = None) -> Union[str, bytes]:
        """
        Converts document content from one format to another.
        
        Args:
            content: Document content to convert
            source_format: Source format (extension without dot)
            target_format: Target format (extension without dot)
            options: Optional conversion parameters
            
        Returns:
            Converted document content
            
        Raises:
            UnsupportedFormatError: If source or target format is not supported
            FormatConversionError: If conversion fails
        """
        if not options:
            options = {}
            
        # Validate formats
        if source_format not in self._import_formats:
            raise UnsupportedFormatError(source_format, "import")
            
        if target_format not in self._export_formats:
            raise UnsupportedFormatError(target_format, "export")
            
        # If source and target are the same, return content unchanged
        if source_format == target_format:
            return content
        
        # Converter key
        converter_key = f"{source_format}:{target_format}"
        
        # Check if we have a direct converter
        if converter_key in self._converters:
            try:
                logger.info(f"Converting from {source_format} to {target_format}")
                return self._converters[converter_key](content, options)
            except Exception as e:
                logger.error(f"Conversion failed from {source_format} to {target_format}: {str(e)}")
                raise FormatConversionError(
                    f"Failed to convert from {source_format} to {target_format}: {str(e)}",
                    source_format, target_format, e
                )
        
        # If no direct converter, try to find a path through intermediate formats
        # For simplicity, we'll only try one intermediate format (HTML is a good hub)
        if source_format != "html" and target_format != "html":
            try:
                # Convert source to HTML
                html_content = self.convert(content, source_format, "html", options)
                
                # Convert HTML to target
                return self.convert(html_content, "html", target_format, options)
            except Exception as e:
                logger.error(f"Multi-step conversion failed from {source_format} to {target_format}: {str(e)}")
                raise FormatConversionError(
                    f"Failed to convert from {source_format} to {target_format} via intermediate format: {str(e)}",
                    source_format, target_format, e
                )
        
        # If we got here, no conversion path was found
        raise FormatConversionError(
            f"No conversion path found from {source_format} to {target_format}",
            source_format, target_format
        )
    
    def register_converter(self, source_format: str, target_format: str, 
                          converter_func: Callable[[Union[str, bytes], Dict[str, Any]], Union[str, bytes]]) -> None:
        """
        Registers a custom converter function for a specific format pair.
        
        Args:
            source_format: Source format (extension without dot)
            target_format: Target format (extension without dot)
            converter_func: Function that converts from source to target format
        """
        # Validate formats
        if not is_valid_document_format(source_format):
            raise UnsupportedFormatError(source_format, "conversion source")
            
        if not is_valid_document_format(target_format):
            raise UnsupportedFormatError(target_format, "conversion target")
        
        # Register the converter
        converter_key = f"{source_format}:{target_format}"
        self._converters[converter_key] = converter_func
        
        # Update the import/export format lists if needed
        if source_format not in self._import_formats:
            self._import_formats.append(source_format)
            
        if target_format not in self._export_formats:
            self._export_formats.append(target_format)
            
        logger.info(f"Registered converter from {source_format} to {target_format}")
    
    def register_detector(self, format_name: str, detector_func: Callable[[Union[str, bytes]], bool]) -> None:
        """
        Registers a custom format detection function.
        
        Args:
            format_name: Format name (extension without dot)
            detector_func: Function that detects if content matches the format
        """
        # Validate format
        if not is_valid_document_format(format_name):
            raise UnsupportedFormatError(format_name, "format detection")
        
        # Register the detector
        self._detectors[format_name] = detector_func
        logger.info(f"Registered detector for {format_name}")
    
    def get_supported_formats(self) -> Dict[str, List[str]]:
        """
        Returns the list of supported import and export formats.
        
        Returns:
            Dictionary with import_formats and export_formats lists
        """
        return {
            "import_formats": self._import_formats.copy(),
            "export_formats": self._export_formats.copy()
        }
    
    def _convert_html_to_markdown(self, content: str, options: Dict[str, Any]) -> str:
        """
        Converts HTML content to Markdown format.
        
        Args:
            content: HTML content
            options: Conversion options
            
        Returns:
            Markdown content
        """
        h2t = html2text.HTML2Text()
        
        # Configure html2text options
        h2t.ignore_links = options.get('ignore_links', False)
        h2t.ignore_images = options.get('ignore_images', False)
        h2t.ignore_tables = options.get('ignore_tables', False)
        h2t.body_width = options.get('body_width', 0)  # 0 means no wrapping
        h2t.protect_links = options.get('protect_links', True)
        h2t.unicode_snob = options.get('unicode_snob', True)  # Use Unicode instead of ASCII
        h2t.escape_snob = options.get('escape_snob', True)  # Don't escape special characters
        
        # Convert HTML to Markdown
        return h2t.handle(content)
    
    def _convert_html_to_text(self, content: str, options: Dict[str, Any]) -> str:
        """
        Converts HTML content to plain text.
        
        Args:
            content: HTML content
            options: Conversion options
            
        Returns:
            Plain text content
        """
        # First convert to Markdown as an intermediate step
        h2t = html2text.HTML2Text()
        
        # Configure html2text for plain text output
        h2t.ignore_links = options.get('ignore_links', True)
        h2t.ignore_images = options.get('ignore_images', True)
        h2t.ignore_tables = options.get('ignore_tables', False)
        h2t.body_width = options.get('body_width', 0)  # 0 means no wrapping
        h2t.unicode_snob = options.get('unicode_snob', True)
        
        # Convert to Markdown
        markdown_text = h2t.handle(content)
        
        # Remove Markdown formatting characters
        plain_text = re.sub(r'[#*_~`\[\]\(\)]', '', markdown_text)
        
        # Normalize whitespace and line breaks
        plain_text = re.sub(r'\n{3,}', '\n\n', plain_text)
        plain_text = re.sub(r'[ \t]+', ' ', plain_text)
        
        return plain_text.strip()
    
    def _convert_markdown_to_html(self, content: str, options: Dict[str, Any]) -> str:
        """
        Converts Markdown content to HTML.
        
        Args:
            content: Markdown content
            options: Conversion options
            
        Returns:
            HTML content
        """
        # Configure markdown extensions
        extensions = options.get('extensions', [
            'markdown.extensions.tables',
            'markdown.extensions.fenced_code',
            'markdown.extensions.footnotes'
        ])
        
        # Convert Markdown to HTML
        html_body = markdown.markdown(content, extensions=extensions)
        
        # Optionally add HTML wrapper
        if options.get('add_wrapper', True):
            html_head = '<!DOCTYPE html>\n<html>\n<head>\n'
            html_head += '<meta charset="utf-8">\n'
            
            # Add title if provided
            if 'title' in options:
                html_head += f'<title>{options["title"]}</title>\n'
            
            # Add stylesheet if provided
            if 'css' in options:
                html_head += f'<style>{options["css"]}</style>\n'
            
            html_head += '</head>\n<body>\n'
            html_foot = '\n</body>\n</html>'
            
            return html_head + html_body + html_foot
        
        return html_body
    
    def _convert_markdown_to_text(self, content: str, options: Dict[str, Any]) -> str:
        """
        Converts Markdown content to plain text.
        
        Args:
            content: Markdown content
            options: Conversion options
            
        Returns:
            Plain text content
        """
        # Remove Markdown formatting characters
        plain_text = content
        
        # Headers
        plain_text = re.sub(r'^#+\s+', '', plain_text, flags=re.MULTILINE)
        
        # Bold, italic, strikethrough
        plain_text = re.sub(r'[*_~]{1,3}(.*?)[*_~]{1,3}', r'\1', plain_text)
        
        # Links
        plain_text = re.sub(r'\[(.*?)\]\(.*?\)', r'\1', plain_text)
        
        # Images
        plain_text = re.sub(r'!\[(.*?)\]\(.*?\)', r'\1', plain_text)
        
        # Code blocks
        plain_text = re.sub(r'```.*?```', '', plain_text, flags=re.DOTALL)
        plain_text = re.sub(r'`(.*?)`', r'\1', plain_text)
        
        # Lists
        plain_text = re.sub(r'^\s*[-*+]\s+', '', plain_text, flags=re.MULTILINE)
        plain_text = re.sub(r'^\s*\d+\.\s+', '', plain_text, flags=re.MULTILINE)
        
        # Blockquotes
        plain_text = re.sub(r'^\s*>\s+', '', plain_text, flags=re.MULTILINE)
        
        # Normalize whitespace and line breaks
        plain_text = re.sub(r'\n{3,}', '\n\n', plain_text)
        plain_text = re.sub(r'[ \t]+', ' ', plain_text)
        
        return plain_text.strip()
    
    def _convert_text_to_html(self, content: str, options: Dict[str, Any]) -> str:
        """
        Converts plain text content to HTML.
        
        Args:
            content: Plain text content
            options: Conversion options
            
        Returns:
            HTML content
        """
        # Escape HTML special characters
        html_content = content.replace('&', '&amp;')
        html_content = html_content.replace('<', '&lt;')
        html_content = html_content.replace('>', '&gt;')
        
        # Convert line breaks to <br> tags
        html_content = html_content.replace('\n', '<br>\n')
        
        # Wrap paragraphs (separated by double line breaks)
        html_paragraphs = []
        for paragraph in re.split(r'<br>\n<br>\n', html_content):
            if paragraph.strip():
                html_paragraphs.append(f'<p>{paragraph}</p>')
        
        html_body = '\n'.join(html_paragraphs)
        
        # Optionally add HTML wrapper
        if options.get('add_wrapper', True):
            html_head = '<!DOCTYPE html>\n<html>\n<head>\n'
            html_head += '<meta charset="utf-8">\n'
            
            # Add title if provided
            if 'title' in options:
                html_head += f'<title>{options["title"]}</title>\n'
            
            # Add stylesheet if provided
            if 'css' in options:
                html_head += f'<style>{options["css"]}</style>\n'
            
            html_head += '</head>\n<body>\n'
            html_foot = '\n</body>\n</html>'
            
            return html_head + html_body + html_foot
        
        return html_body
    
    def _convert_docx_to_html(self, content: bytes, options: Dict[str, Any]) -> str:
        """
        Converts DOCX content to HTML.
        
        Args:
            content: DOCX content as bytes
            options: Conversion options
            
        Returns:
            HTML content
        """
        # Create BytesIO object for mammoth to read
        docx_file = io.BytesIO(content)
        
        # Set up conversion options
        convert_options = {}
        
        # Convert DOCX to HTML
        result = mammoth.convert_to_html(docx_file, **convert_options)
        html_body = result.value
        
        # Log warnings if any
        for warning in result.messages:
            logger.warning(f"DOCX conversion warning: {warning}")
        
        # Optionally add HTML wrapper
        if options.get('add_wrapper', True):
            html_head = '<!DOCTYPE html>\n<html>\n<head>\n'
            html_head += '<meta charset="utf-8">\n'
            
            # Add title if provided
            if 'title' in options:
                html_head += f'<title>{options["title"]}</title>\n'
            
            # Add stylesheet if provided
            if 'css' in options:
                html_head += f'<style>{options["css"]}</style>\n'
            else:
                # Add default styles for DOCX conversion
                html_head += '<style>\n'
                html_head += 'body { font-family: Arial, sans-serif; line-height: 1.6; }\n'
                html_head += 'h1, h2, h3, h4, h5, h6 { margin-top: 1.2em; margin-bottom: 0.8em; }\n'
                html_head += 'p { margin: 0.8em 0; }\n'
                html_head += 'table { border-collapse: collapse; margin: 1em 0; }\n'
                html_head += 'td, th { border: 1px solid #ddd; padding: 0.5em; }\n'
                html_head += '</style>\n'
            
            html_head += '</head>\n<body>\n'
            html_foot = '\n</body>\n</html>'
            
            return html_head + html_body + html_foot
        
        return html_body
    
    def _convert_rtf_to_text(self, content: Union[str, bytes], options: Dict[str, Any]) -> str:
        """
        Converts RTF content to plain text.
        
        Args:
            content: RTF content
            options: Conversion options
            
        Returns:
            Plain text content
        """
        # Ensure content is a string
        if isinstance(content, bytes):
            content_str = content.decode('utf-8', errors='ignore')
        else:
            content_str = content
        
        # Convert RTF to plain text using striprtf
        plain_text = striprtf.rtf_to_text(content_str)
        
        # Clean up any conversion artifacts
        plain_text = re.sub(r'\r\n', '\n', plain_text)  # Normalize line endings
        plain_text = re.sub(r'\n{3,}', '\n\n', plain_text)  # Remove excessive line breaks
        
        return plain_text.strip()
    
    def _convert_html_to_pdf(self, content: str, options: Dict[str, Any]) -> bytes:
        """
        Converts HTML content to PDF.
        
        Args:
            content: HTML content
            options: Conversion options
            
        Returns:
            PDF content as bytes
        """
        # Create a BytesIO object to store the PDF output
        pdf_file = io.BytesIO()
        
        # Set up conversion options
        pdf_options = {
            'page_size': options.get('page_size', 'A4'),
            'margin_top': options.get('margin_top', '1cm'),
            'margin_right': options.get('margin_right', '1cm'),
            'margin_bottom': options.get('margin_bottom', '1cm'),
            'margin_left': options.get('margin_left', '1cm')
        }
        
        # Add CSS if provided
        if 'css' in options:
            # WeasyPrint can use a stylesheet
            stylesheet = weasyprint.CSS(string=options['css'])
            # Generate PDF with the custom stylesheet
            weasyprint.HTML(string=content).write_pdf(
                pdf_file, 
                stylesheets=[stylesheet],
                presentational_hints=True
            )
        else:
            # Generate PDF with default styling
            weasyprint.HTML(string=content).write_pdf(
                pdf_file,
                presentational_hints=True
            )
        
        # Reset file pointer to beginning
        pdf_file.seek(0)
        
        # Return PDF content as bytes
        return pdf_file.getvalue()
    
    def _convert_html_to_docx(self, content: str, options: Dict[str, Any]) -> bytes:
        """
        Converts HTML content to DOCX format.
        
        Args:
            content: HTML content
            options: Conversion options
            
        Returns:
            DOCX content as bytes
        """
        # Create a new docx document
        document = docx.Document()
        
        # Parse HTML
        # This is a simplified implementation; a complete solution would need 
        # to handle a full range of HTML elements
        
        # Remove doctype, html, head, and body tags to get just the content
        content_body = re.sub(r'<!DOCTYPE.*?>', '', content, flags=re.IGNORECASE | re.DOTALL)
        content_body = re.sub(r'<html.*?>.*?<body.*?>', '', content_body, flags=re.IGNORECASE | re.DOTALL)
        content_body = re.sub(r'</body>.*?</html>', '', content_body, flags=re.IGNORECASE | re.DOTALL)
        
        # Split by heading and paragraph tags
        elements = re.split(r'(<h[1-6].*?>.*?</h[1-6]>|<p.*?>.*?</p>)', content_body, flags=re.IGNORECASE | re.DOTALL)
        
        for element in elements:
            if not element.strip():
                continue
                
            # Check if it's a heading
            heading_match = re.match(r'<h([1-6]).*?>(.*?)</h\1>', element, re.IGNORECASE | re.DOTALL)
            if heading_match:
                level = int(heading_match.group(1))
                text = re.sub(r'<.*?>', '', heading_match.group(2)).strip()
                
                if level == 1:
                    document.add_heading(text, 0)  # Title
                else:
                    document.add_heading(text, level)
                continue
            
            # Check if it's a paragraph
            para_match = re.match(r'<p.*?>(.*?)</p>', element, re.IGNORECASE | re.DOTALL)
            if para_match:
                text = re.sub(r'<.*?>', '', para_match.group(1)).strip()
                if text:
                    document.add_paragraph(text)
                continue
            
            # If we get here, it's not a heading or paragraph, add as plain text
            text = re.sub(r'<.*?>', '', element).strip()
            if text:
                document.add_paragraph(text)
        
        # Save the document to a BytesIO object
        docx_file = io.BytesIO()
        document.save(docx_file)
        
        # Reset file pointer to beginning
        docx_file.seek(0)
        
        # Return DOCX content as bytes
        return docx_file.getvalue()
    
    def _detect_html(self, content: Union[str, bytes]) -> bool:
        """
        Detects if content is in HTML format.
        
        Args:
            content: Content to analyze
            
        Returns:
            True if content appears to be HTML
        """
        # Convert bytes to string if needed
        if isinstance(content, bytes):
            try:
                content_str = content.decode('utf-8')
            except UnicodeDecodeError:
                return False  # Not text content
        else:
            content_str = content
            
        # Check for HTML indicators
        html_patterns = [
            r'<!DOCTYPE\s+html',
            r'<html.*?>',
            r'<head.*?>',
            r'<body.*?>',
            r'<div.*?>',
            r'<p.*?>',
            r'<h[1-6].*?>'
        ]
        
        for pattern in html_patterns:
            if re.search(pattern, content_str, re.IGNORECASE | re.DOTALL):
                return True
        
        # Check for multiple HTML tags
        tag_pattern = r'<[a-z][a-z0-9]*(\s+[^>]*)?>'
        tag_matches = re.findall(tag_pattern, content_str, re.IGNORECASE)
        if len(tag_matches) > 5:  # Arbitrary threshold for multiple tags
            return True
        
        return False
    
    def _detect_markdown(self, content: Union[str, bytes]) -> bool:
        """
        Detects if content is in Markdown format.
        
        Args:
            content: Content to analyze
            
        Returns:
            True if content appears to be Markdown
        """
        # Convert bytes to string if needed
        if isinstance(content, bytes):
            try:
                content_str = content.decode('utf-8')
            except UnicodeDecodeError:
                return False  # Not text content
        else:
            content_str = content
        
        # Check for Markdown indicators
        md_patterns = [
            r'^#{1,6}\s+.+$',  # Headers
            r'^-{3,}$',         # Horizontal rules
            r'^\*\s+.+$',       # Unordered lists
            r'^\d+\.\s+.+$',    # Ordered lists
            r'\[.+\]\(.+\)',    # Links
            r'!\[.+\]\(.+\)',   # Images
            r'`{1,3}',          # Code blocks
            r'\*\*.+\*\*',      # Bold
            r'_.+_',            # Italics
        ]
        
        md_pattern_count = 0
        for pattern in md_patterns:
            matches = re.findall(pattern, content_str, re.MULTILINE)
            md_pattern_count += len(matches)
        
        # If we found multiple Markdown patterns, it's likely Markdown
        # But need to exclude content that's also valid HTML
        if md_pattern_count >= 3 and not self._detect_html(content_str):
            return True
        
        return False
    
    def _detect_docx(self, content: Union[str, bytes]) -> bool:
        """
        Detects if content is in DOCX format.
        
        Args:
            content: Content to analyze
            
        Returns:
            True if content appears to be DOCX
        """
        # DOCX is a binary format, so we need bytes
        if isinstance(content, str):
            return False
        
        # Check for the DOCX file signature (PK zip format + specific content types)
        # DOCX files are ZIP archives with specific structure
        if len(content) < 4:
            return False
        
        # Check for PK zip header
        if content[:4] != b'PK\x03\x04':
            return False
        
        # Look for common DOCX contents
        docx_markers = [
            b'word/',
            b'docProps/',
            b'content-types',
            b'document.xml'
        ]
        
        for marker in docx_markers:
            if marker in content:
                return True
        
        return False
    
    def _detect_rtf(self, content: Union[str, bytes]) -> bool:
        """
        Detects if content is in RTF format.
        
        Args:
            content: Content to analyze
            
        Returns:
            True if content appears to be RTF
        """
        # Convert bytes to string if needed
        if isinstance(content, bytes):
            try:
                # Just check the first chunk for the signature
                content_start = content[:100].decode('utf-8', errors='ignore')
            except Exception:
                return False
        else:
            content_start = content[:100] if len(content) > 100 else content
        
        # RTF files always start with {\\rtf
        return content_start.startswith('{\\rtf')